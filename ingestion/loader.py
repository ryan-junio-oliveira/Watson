import logging
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import List, Optional

import docx
import pypdf

try:
    from PIL import Image
    from pdf2image import convert_from_path
    import pytesseract

    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False


@dataclass
class LoadedDocument:
    content: str
    filepath: str
    filename: str
    file_type: str
    modified_at: str
    file_size: int


class DocumentLoader:
    SUPPORTED_EXTENSIONS = {
        ".pdf", ".docx", ".txt", ".md",
        ".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif",
    }

    def __init__(self, logger: Optional[logging.Logger] = None):
        self.logger = logger

    def load(self, directory: str) -> List[LoadedDocument]:
        dir_path = Path(directory)
        if not dir_path.exists():
            raise FileNotFoundError(f"Directory not found: {directory}")
        if not dir_path.is_dir():
            raise NotADirectoryError(f"Path is not a directory: {directory}")

        documents: List[LoadedDocument] = []
        for filepath in sorted(dir_path.rglob("*")):
            if not filepath.is_file():
                continue
            ext = filepath.suffix.lower()
            if ext not in self.SUPPORTED_EXTENSIONS:
                self._log_warning(f"Unsupported file type skipped: {filepath}")
                continue

            try:
                doc = self._load_single(filepath)
                documents.append(doc)
                self._log_info(f"Loaded document: {filepath.name}")
            except Exception as e:
                self._log_error(f"Failed to load {filepath}: {e}")
                continue

        self._log_info(f"Loaded {len(documents)} documents from {directory}")
        return documents

    def _load_single(self, filepath: Path) -> LoadedDocument:
        ext = filepath.suffix.lower()
        loaders = {
            ".pdf": self._load_pdf,
            ".docx": self._load_docx,
            ".md": self._load_txt,
            ".txt": self._load_txt,
            ".jpg": self._load_image,
            ".jpeg": self._load_image,
            ".png": self._load_image,
            ".bmp": self._load_image,
            ".tiff": self._load_image,
            ".tif": self._load_image,
        }
        loader = loaders.get(ext)
        if loader is None:
            raise ValueError(f"Unsupported file type: {ext}")

        content = loader(filepath)
        stat = filepath.stat()
        modified_at = datetime.fromtimestamp(stat.st_mtime).isoformat()

        return LoadedDocument(
            content=content,
            filepath=str(filepath.absolute()),
            filename=filepath.name,
            file_type=ext,
            modified_at=modified_at,
            file_size=stat.st_size,
        )

    def _load_pdf(self, filepath: Path) -> str:
        text_parts: List[str] = []
        with open(filepath, "rb") as f:
            reader = pypdf.PdfReader(f)
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

        extracted = "\n".join(text_parts)

        if extracted.strip():
            return extracted

        self._log_warning(
            f"No text extracted by pypdf from '{filepath.name}', "
            f"attempting OCR fallback..."
        )

        if not OCR_AVAILABLE:
            self._log_warning(
                "OCR not available. Install pdf2image and pytesseract: "
                "pip install pdf2image pytesseract"
            )
            return ""

        try:
            images = convert_from_path(str(filepath), dpi=300)
            ocr_text_parts: List[str] = []
            for i, image in enumerate(images):
                text = pytesseract.image_to_string(image, lang="por+eng")
                if text.strip():
                    ocr_text_parts.append(text)
                self._log_info(
                    f"OCR page {i + 1}/{len(images)}: "
                    f"{len(text.strip())} chars extracted"
                )
            ocr_result = "\n".join(ocr_text_parts)
            if ocr_result.strip():
                self._log_info(
                    f"OCR extracted {len(ocr_result.strip())} chars from "
                    f"'{filepath.name}'"
                )
            else:
                self._log_warning(
                    f"OCR also produced no text from '{filepath.name}'"
                )
            return ocr_result
        except Exception as e:
            self._log_error(f"OCR failed for '{filepath.name}': {e}")
            return ""

    def _load_docx(self, filepath: Path) -> str:
        doc = docx.Document(str(filepath))
        text_parts = [p.text for p in doc.paragraphs if p.text]
        return "\n".join(text_parts)

    def _load_txt(self, filepath: Path) -> str:
        return filepath.read_text(encoding="utf-8", errors="replace")

    def _load_image(self, filepath: Path) -> str:
        if not OCR_AVAILABLE:
            self._log_warning(
                f"Cannot extract text from image '{filepath.name}': "
                "pytesseract not installed"
            )
            return ""

        try:
            image = Image.open(str(filepath))
            text = pytesseract.image_to_string(image, lang="por+eng")
            result = text.strip()
            if result:
                self._log_info(
                    f"OCR extracted {len(result)} chars from image "
                    f"'{filepath.name}'"
                )
            else:
                self._log_warning(
                    f"No text found in image '{filepath.name}'"
                )
            return result
        except Exception as e:
            self._log_error(f"Image OCR failed for '{filepath.name}': {e}")
            return ""

    def _log_info(self, message: str) -> None:
        if self.logger:
            self.logger.info(message)

    def _log_warning(self, message: str) -> None:
        if self.logger:
            self.logger.warning(message)

    def _log_error(self, message: str) -> None:
        if self.logger:
            self.logger.error(message)
