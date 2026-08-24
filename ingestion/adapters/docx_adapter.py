"""Adapter de DOCX: headings, parágrafos, tabelas e imagens (§2)."""

from __future__ import annotations

import logging
from pathlib import Path
from typing import List, Optional

import docx

from ingestion.adapters.base import SourceAdapter
from ingestion.models import ImageRef, LoadedDocument, Section, Table, sha256_text

_HEADING_STYLES = {
    "Heading 1": 1,
    "Heading 2": 2,
    "Heading 3": 3,
    "Heading 4": 4,
    "Heading 5": 5,
    "Heading 6": 6,
}


class DocxAdapter(SourceAdapter):
    source_type = "docx"
    supported_extensions = {".docx"}

    def __init__(self, logger: Optional[logging.Logger] = None):
        self.logger = logger

    def extract(self, filepath: Path) -> LoadedDocument:
        doc = docx.Document(str(filepath))
        sections: List[Section] = []
        tables: List[Table] = []
        paragraphs: List[str] = []
        images: List[ImageRef] = []

        current_section = ""

        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ""
            level = _HEADING_STYLES.get(style_name, 0)
            text = para.text.strip()
            if not text and level == 0:
                continue
            if level > 0:
                sections.append(Section(heading=text, level=level))
                current_section = text
                paragraphs.append(f"{'#' * level} {text}")
            else:
                paragraphs.append(text)

        for table_index, table in enumerate(doc.tables):
            headers: List[str] = []
            rows: List[List[str]] = []
            for row_index, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                if row_index == 0:
                    headers = cells
                rows.append(cells)
            tables.append(
                Table(
                    section=current_section,
                    markdown=_table_to_markdown(headers, rows),
                    headers=headers,
                    rows=rows,
                )
            )
            paragraphs.append(f"\n[Tabela {table_index + 1}]\n{_table_to_markdown(headers, rows)}")

        for i, rel in enumerate(_iter_docx_images(doc)):
            images.append(ImageRef(image_id=f"img_{i}", section=current_section))

        content = "\n".join(paragraphs)
        return LoadedDocument(
            content=content,
            filepath=str(filepath),
            filename=filepath.name,
            file_type=".docx",
            modified_at="",
            file_size=0,
            source_type=self.source_type,
            source_id=str(filepath),
            metadata={"paragraphs": len(paragraphs), "tables": len(tables)},
            sections=sections,
            tables=tables,
            images=images,
            content_hash=sha256_text(content),
        )


def _iter_docx_images(doc):
    from docx.parts.image import ImagePart

    for rel in doc.part.rels.values():
        if isinstance(rel.target_part, ImagePart):
            yield rel.target_part


def _table_to_markdown(headers: List[str], rows: List[List[str]]) -> str:
    lines: List[str] = []
    if headers:
        lines.append("| " + " | ".join(headers) + " |")
        lines.append("|" + "---|" * len(headers))
    for row in rows:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)
