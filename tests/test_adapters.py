import csv as csv_module
from pathlib import Path

import pymupdf
import pytest

from ingestion.adapters.csv_adapter import CsvAdapter
from ingestion.adapters.docx_adapter import DocxAdapter
from ingestion.adapters.image_adapter import ImageAdapter
from ingestion.adapters.pdf_adapter import PdfAdapter
from ingestion.adapters.registry import build_default_registry
from ingestion.adapters.text_adapter import TextAdapter
from ingestion.adapters.xlsx_adapter import XlsxAdapter
from ingestion.models import LoadedDocument


def make_simple_pdf(path: Path, text: str = "Conteudo de teste PDF.") -> Path:
    doc = pymupdf.open()
    page = doc.new_page()
    page.insert_text((72, 72), text)
    doc.save(str(path))
    doc.close()
    return path


class TestTextAdapter:
    def test_txt_content(self, tmp_path):
        f = tmp_path / "doc.txt"
        f.write_text("linha um\nlinha dois\n", encoding="utf-8")
        doc = TextAdapter().extract(f)
        assert isinstance(doc, LoadedDocument)
        assert "linha um" in doc.content
        assert doc.source_type == "text"

    def test_md_headings(self, tmp_path):
        f = tmp_path / "doc.md"
        f.write_text("# Titulo\n\n# Segundo\n\ntexto\n", encoding="utf-8")
        doc = TextAdapter().extract(f)
        headings = [s.heading for s in doc.sections]
        assert "Titulo" in headings
        assert "Segundo" in headings


class TestCsvAdapter:
    def test_csv_structure(self, tmp_path):
        f = tmp_path / "data.csv"
        with open(f, "w", newline="", encoding="utf-8") as fh:
            writer = csv_module.writer(fh)
            writer.writerow(["id", "nome"])
            writer.writerow(["1", "HP MODELO-X"])
        doc = CsvAdapter().extract(f)
        assert len(doc.tables) == 1
        assert doc.tables[0].headers == ["id", "nome"]
        assert doc.tables[0].rows == [["1", "HP MODELO-X"]]
        assert "1 | HP MODELO-X" in doc.content


class TestXlsxAdapter:
    def test_xlsx_sheet(self, tmp_path):
        import openpyxl

        f = tmp_path / "book.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Impressoras"
        ws.append(["modelo", "tipo"])
        ws.append(["MODELO-X", "printer"])
        wb.save(f)
        wb.close()

        doc = XlsxAdapter().extract(f)
        assert doc.source_type == "xlsx"
        assert any(s.heading == "Planilha: Impressoras" for s in doc.sections)
        assert len(doc.tables) == 1
        assert doc.tables[0].headers == ["modelo", "tipo"]
        assert "Impressoras" in doc.content


class TestDocxAdapter:
    def test_docx_headings_and_tables(self, tmp_path):
        import docx

        f = tmp_path / "doc.docx"
        d = docx.Document()
        d.add_heading("Manual", level=1)
        d.add_paragraph("Passo para resolver.")
        table = d.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "codigo"
        table.cell(0, 1).text = "descricao"
        table.cell(1, 0).text = "E123"
        table.cell(1, 1).text = "papel preso"
        d.save(f)

        doc = DocxAdapter().extract(f)
        assert doc.source_type == "docx"
        assert any(s.heading == "Manual" for s in doc.sections)
        assert "Passo para resolver" in doc.content
        assert len(doc.tables) >= 1


class TestPdfAdapter:
    def test_simple_pdf_text(self, tmp_path):
        pdf_path = make_simple_pdf(tmp_path / "simple.pdf", "Instrucao de teste E123.")
        doc = PdfAdapter().extract(pdf_path)
        assert doc.source_type == "pdf"
        assert len(doc.pages) == 1
        assert doc.metadata["pages"] == 1
        assert "E123" in doc.content

    def test_scanned_page_triggers_ocr_gracefully(self, tmp_path):
        pdf_path = tmp_path / "scanned.pdf"
        doc = pymupdf.open()
        doc.new_page()
        doc.save(str(pdf_path))
        doc.close()

        adapter = PdfAdapter()
        result = adapter.extract(pdf_path)
        # página sem texto é candidata a OCR; se OCR indisponível, não quebra
        assert result.pages[0].ocr is True
        assert result.metadata["ocr_pages"] == 1

    def test_needs_ocr_heuristic(self):
        adapter = PdfAdapter()
        assert adapter._needs_ocr("") is True
        assert adapter._needs_ocr("x") is True
        assert adapter._needs_ocr("texto com conteudo real de documento") is False

    def test_section_extraction(self):
        adapter = PdfAdapter()
        md = {0: "# Troubleshooting\n\ntexto\n\n## Error Codes\n\nmais texto\n"}
        sections = adapter._extract_sections(md, [])
        assert [s.heading for s in sections] == ["Troubleshooting", "Error Codes"]
        assert sections[0].level == 1
        assert sections[1].level == 2

    def test_table_extraction(self):
        adapter = PdfAdapter()
        md = {
            0: (
                "| codigo | descricao |\n"
                "|---|---|\n"
                "| E123 | papel preso |\n"
                "| E124 | atolamento |\n"
            )
        }
        tables = adapter._extract_tables(md, [])
        assert len(tables) == 1
        assert tables[0].headers == ["codigo", "descricao"]
        assert tables[0].rows == [["E123", "papel preso"], ["E124", "atolamento"]]


class TestImageAdapter:
    def test_image_extract_no_ocr_binary(self, tmp_path):
        from PIL import Image

        f = tmp_path / "img.png"
        Image.new("RGB", (200, 100), color="white").save(f)

        doc = ImageAdapter().extract(f)
        assert doc.source_type == "image"
        assert doc.metadata["image_format"] == "PNG"
        assert doc.metadata["image_width"] == 200


class TestRegistry:
    def test_default_registry_supported_extensions(self):
        registry = build_default_registry()
        assert ".pdf" in registry.supported_extensions
        assert ".docx" in registry.supported_extensions
        assert ".csv" in registry.supported_extensions
        assert ".xlsx" in registry.supported_extensions
        assert ".png" in registry.supported_extensions
        assert ".txt" in registry.supported_extensions

    def test_registry_dispatches(self, tmp_path):
        from ingestion.loader import DocumentLoader

        f = tmp_path / "doc.md"
        f.write_text("# Titulo\n\ntexto\n", encoding="utf-8")
        loader = DocumentLoader()
        doc = loader._load_single(f)
        assert doc.source_type == "text"
        assert doc.modified_at
        assert doc.file_size == f.stat().st_size

class TestPdfImagePersistence:
    def test_images_saved_to_disk(self, tmp_path):
        from PIL import Image
        import io
        import pymupdf

        buf = io.BytesIO()
        Image.new("RGB", (400, 300), "blue").save(buf, format="PNG")
        png_bytes = buf.getvalue()

        pdf_path = tmp_path / "with_img.pdf"
        doc = pymupdf.open()
        page = doc.new_page(width=595, height=842)
        page.insert_text((72, 72), "texto da pagina")
        page.insert_image(pymupdf.Rect(100, 100, 500, 400), stream=png_bytes)
        doc.save(str(pdf_path))
        doc.close()

        image_dir = tmp_path / "images"
        adapter = PdfAdapter(image_dir=str(image_dir))
        result = adapter.extract(pdf_path)
        assert len(result.images) >= 1
        saved = [i for i in result.images if i.storage_path]
        assert len(saved) >= 1
        assert Path(saved[0].storage_path).exists()


class TestLoaderIdentity:
    def test_manufacturer_inferred_from_filename(self, tmp_path):
        from ingestion.loader import DocumentLoader

        f = tmp_path / "HP_MODELO-X_manual.txt"
        f.write_text("texto", encoding="utf-8")
        loader = DocumentLoader()
        doc = loader._load_single(f)
        assert doc.metadata.get("manufacturer") == "HP"
        assert "MODELO-X" in doc.metadata.get("model", "")
